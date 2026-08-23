from io import BytesIO

from docx import Document as DocxDocument
from langchain_core.documents import Document

from app.utils.logger import get_logger

logger = get_logger(__name__)


class DOCXLoader:
    """Loads a Word document from in-memory content."""

    def __init__(self, file_content: bytes, filename: str):
        self.file_content = file_content
        self.filename = filename

    def load(self):
        """Load the Word document."""

        logger.info("Loading Word document: %s", self.filename)

        try:
            docx = DocxDocument(BytesIO(self.file_content))

            text = "\n".join(
                paragraph.text
                for paragraph in docx.paragraphs
            )

            documents = [
                Document(
                    page_content=text,
                    metadata={"source": self.filename},
                )
            ]

            logger.info(
                "Successfully loaded Word document '%s'.",
                self.filename,
            )

            return documents

        except Exception:
            logger.exception(
                "Failed to load Word document: %s",
                self.filename,
            )
            raise